from io import BytesIO

from docx import Document
from openpyxl import Workbook

from medreg.modules.documents.parser import ControlledDocumentParser


def test_html_parser_keeps_visible_text() -> None:
    parser = ControlledDocumentParser()

    result = parser.extract(
        "regulation.html",
        "<html><body><h1>管理办法</h1><p>适用于境内注册。</p></body></html>".encode(),
    )

    assert result == "管理办法\n适用于境内注册。"


def test_html_parser_drops_script_and_style_content() -> None:
    parser = ControlledDocumentParser()

    result = parser.extract(
        "regulation.html",
        (
            "<style>.title { color: red; }</style>"
            "<script>window.secret = 'ignore';</script>"
            "<h1>医疗器械管理办法</h1>"
        ).encode(),
    )

    assert result == "医疗器械管理办法"


def test_docx_parser_extracts_paragraphs_and_tables() -> None:
    document = Document()
    document.add_paragraph("产品技术要求")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "型号"
    table.cell(0, 1).text = "MR-01"
    buffer = BytesIO()
    document.save(buffer)

    result = ControlledDocumentParser().extract("requirements.docx", buffer.getvalue())

    assert "产品技术要求" in result
    assert "型号 MR-01" in result


def test_xlsx_parser_extracts_sheets_and_structured_tables() -> None:
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "申报资料矩阵"
    worksheet.append(["资料类别", "状态", "责任人"])
    worksheet.append(["风险分析", "已接受", "刘凯旗"])
    worksheet.append(["临床评价", "待补充", "张法规"])
    buffer = BytesIO()
    workbook.save(buffer)

    parsed = ControlledDocumentParser().parse("dossier-matrix.xlsx", buffer.getvalue())

    assert "工作表：申报资料矩阵" in parsed.text
    assert len(parsed.tables) == 1
    assert parsed.tables[0].headers == ("资料类别", "状态", "责任人")
    assert parsed.tables[0].rows[1] == ("临床评价", "待补充", "张法规")
