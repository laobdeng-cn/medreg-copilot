import re
from dataclasses import dataclass
from html.parser import HTMLParser
from io import BytesIO
from pathlib import Path
from typing import Protocol

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader


@dataclass(frozen=True)
class ParsedTable:
    ordinal: int
    title: str
    sheet_name: str | None
    headers: tuple[str, ...]
    rows: tuple[tuple[str, ...], ...]
    source_locator: str


@dataclass(frozen=True)
class ParsedDocument:
    text: str
    tables: tuple[ParsedTable, ...] = ()


class DocumentParser(Protocol):
    version: str

    def extract(self, file_name: str, data: bytes) -> str: ...

    def parse(self, file_name: str, data: bytes) -> ParsedDocument: ...


class _HTMLTextExtractor(HTMLParser):
    ignored_tags = {"script", "style", "noscript", "svg"}

    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []
        self._ignored_depth = 0

    def handle_starttag(
        self, tag: str, attrs: list[tuple[str, str | None]]
    ) -> None:
        del attrs
        if tag.lower() in self.ignored_tags:
            self._ignored_depth += 1

    def handle_endtag(self, tag: str) -> None:
        if tag.lower() in self.ignored_tags and self._ignored_depth:
            self._ignored_depth -= 1

    def handle_data(self, data: str) -> None:
        if self._ignored_depth:
            return
        value = data.strip()
        if value:
            self.parts.append(value)


class ControlledDocumentParser:
    version = "controlled-parser-v3"

    def extract(self, file_name: str, data: bytes) -> str:
        return self.parse(file_name, data).text

    def parse(self, file_name: str, data: bytes) -> ParsedDocument:
        suffix = Path(file_name).suffix.lower()
        if suffix == ".pdf":
            text = self._extract_pdf(data)
            tables: tuple[ParsedTable, ...] = ()
        elif suffix == ".docx":
            text, tables = self._extract_docx(data)
        elif suffix == ".xlsx":
            text, tables = self._extract_xlsx(data)
        elif suffix in {".txt", ".md"}:
            text = data.decode("utf-8-sig")
            tables = ()
        elif suffix in {".html", ".htm"}:
            text = self._extract_html(data)
            tables = ()
        else:
            raise ValueError(f"Unsupported document type: {suffix or 'unknown'}")

        normalized = self._normalize(text)
        if not normalized:
            raise ValueError("No extractable text was found in the document")
        return ParsedDocument(text=normalized, tables=tables)

    @staticmethod
    def _extract_pdf(data: bytes) -> str:
        reader = PdfReader(BytesIO(data))
        return "\n\n".join(page.extract_text() or "" for page in reader.pages)

    @staticmethod
    def _extract_docx(data: bytes) -> tuple[str, tuple[ParsedTable, ...]]:
        document = Document(BytesIO(data))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        tables: list[ParsedTable] = []
        for ordinal, table in enumerate(document.tables, start=1):
            rows = tuple(
                tuple(cell.text.strip() for cell in row.cells) for row in table.rows
            )
            if not rows:
                continue
            headers, body = ControlledDocumentParser._split_table_rows(rows)
            tables.append(
                ParsedTable(
                    ordinal=ordinal,
                    title=f"Word 表格 {ordinal}",
                    sheet_name=None,
                    headers=headers,
                    rows=body,
                    source_locator=f"table:{ordinal}",
                )
            )
            paragraphs.append(f"表格：Word 表格 {ordinal}")
            paragraphs.extend("\t".join(row) for row in rows)
        return "\n".join(paragraphs), tuple(tables)

    @staticmethod
    def _extract_xlsx(data: bytes) -> tuple[str, tuple[ParsedTable, ...]]:
        workbook = load_workbook(BytesIO(data), read_only=True, data_only=False)
        text_parts: list[str] = []
        tables: list[ParsedTable] = []
        ordinal = 0
        for worksheet in workbook.worksheets:
            rows = tuple(
                tuple(ControlledDocumentParser._cell_text(value) for value in row)
                for row in worksheet.iter_rows(values_only=True)
            )
            non_empty = tuple(row for row in rows if any(value for value in row))
            if not non_empty:
                continue
            ordinal += 1
            width = max(len(row) for row in non_empty)
            normalized_rows = tuple(
                row + ("",) * (width - len(row)) for row in non_empty
            )
            headers, body = ControlledDocumentParser._split_table_rows(normalized_rows)
            tables.append(
                ParsedTable(
                    ordinal=ordinal,
                    title=worksheet.title,
                    sheet_name=worksheet.title,
                    headers=headers,
                    rows=body,
                    source_locator=f"sheet:{worksheet.title}!A1",
                )
            )
            text_parts.append(f"工作表：{worksheet.title}")
            text_parts.extend("\t".join(row) for row in normalized_rows)
        workbook.close()
        return "\n".join(text_parts), tuple(tables)

    @staticmethod
    def _split_table_rows(
        rows: tuple[tuple[str, ...], ...],
    ) -> tuple[tuple[str, ...], tuple[tuple[str, ...], ...]]:
        first = rows[0]
        headers = tuple(
            value or f"列 {index}" for index, value in enumerate(first, start=1)
        )
        return headers, rows[1:]

    @staticmethod
    def _cell_text(value: object) -> str:
        if value is None:
            return ""
        return str(value).strip()

    @staticmethod
    def _extract_html(data: bytes) -> str:
        parser = _HTMLTextExtractor()
        parser.feed(data.decode("utf-8-sig"))
        return "\n".join(parser.parts)

    @staticmethod
    def _normalize(text: str) -> str:
        lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.splitlines()]
        return "\n".join(line for line in lines if line)
